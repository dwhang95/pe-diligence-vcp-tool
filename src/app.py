"""
app.py — Streamlit UI for the PE Ops Tool Suite.

Modules:
  Tab 1 — Deal Diligence Brief (Module 1)
  Tab 2 — Value Creation Planner (Module 2)

Run with:
    streamlit run src/app.py
"""

import asyncio
import io
import re
import sys
import threading
import queue
from datetime import datetime
from pathlib import Path

import streamlit as st

# ---------------------------------------------------------------------------
# Path setup — ensure project root is importable
# ---------------------------------------------------------------------------
BASE_DIR = Path(__file__).parent.parent
sys.path.insert(0, str(BASE_DIR))
sys.path.insert(0, str(BASE_DIR / "src"))

# ---------------------------------------------------------------------------
# Page config — must be the first Streamlit call
# ---------------------------------------------------------------------------
st.set_page_config(
    page_title="PE Ops Tool Suite",
    page_icon="📋",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ---------------------------------------------------------------------------
# Theme system — dark (default) and light
# Uses __VARNAME__ placeholders so CSS braces don't need escaping.
# ---------------------------------------------------------------------------

_DARK = dict(
    BG_MAIN         = "#0e1117",
    BG_SIDEBAR      = "#0b0e16",
    BG_PANEL        = "#161b27",
    BG_INPUT        = "#1c2030",
    BG_TABLE_HDR    = "#1f2535",
    BG_BTN_SEC      = "#1c2030",
    BG_EXPANDER     = "#161b27",
    BG_ERROR        = "#1e1424",
    BORDER          = "#2a2d35",
    BORDER_INPUT    = "#2f3347",
    BORDER_ROW      = "#242836",
    TEXT_MAIN       = "#e8e8e8",
    TEXT_BODY       = "#c5c9dc",
    TEXT_LABEL      = "#b0b5c9",
    TEXT_MUTED      = "#8b8fa8",
    TEXT_HEADING    = "#f0f0f0",
    TEXT_H3         = "#d4d8ea",
    TEXT_CODE       = "#b8c0d8",
    TEXT_BQ         = "#8b8fa8",
    ACCENT          = "#c9a84c",
    ACCENT_HOVER    = "#e0be6a",
    ACCENT_FOCUS    = "rgba(201,168,76,0.18)",
    ACCENT_BG_HOVER = "rgba(201,168,76,0.1)",
    BTN_TEXT        = "#0e1117",
    TAB_INACTIVE    = "#6b7090",
    ERR_BG          = "#1e1424",
    ERR_BORDER      = "#5c2d3f",
    ERR_TEXT        = "#f0a0a8",
    LOCK_COLOR      = "#5a5e75",
    TIER_STD_BG     = "#1f2535",
    TIER_STD_TEXT   = "#8b8fa8",
    TIER_STD_BORDER = "#2f3347",
    TIER_PREM_BG    = "rgba(201,168,76,0.12)",
    TIER_PREM_TEXT  = "#c9a84c",
    TIER_PREM_BORDER= "#c9a84c",
    STATS_COLOR     = "#6b7090",
    STATS_STRONG    = "#8b8fa8",
    META_COLOR      = "#8b8fa8",
    META_STRONG     = "#c9a84c",
    EMPTY_BORDER    = "#2a2d35",
    EMPTY_COLOR     = "#3d4259",
    EMPTY_STRONG    = "#4a5070",
    EMPTY_SUB       = "#363b54",
    TOGGLE_BG       = "rgba(255,255,255,0.06)",
    TOGGLE_LABEL    = "#8b8fa8",
)

_LIGHT = dict(
    BG_MAIN         = "#ffffff",
    BG_SIDEBAR      = "#f0f2f8",
    BG_PANEL        = "#f5f5f5",
    BG_INPUT        = "#ffffff",
    BG_TABLE_HDR    = "#eaecf4",
    BG_BTN_SEC      = "#eaecf4",
    BG_EXPANDER     = "#f5f5f5",
    BG_ERROR        = "#fff0f2",
    BORDER          = "#dde1ef",
    BORDER_INPUT    = "#c0c8e0",
    BORDER_ROW      = "#e4e8f2",
    TEXT_MAIN       = "#1a2744",
    TEXT_BODY       = "#333333",
    TEXT_LABEL      = "#1a2744",
    TEXT_MUTED      = "#555e7a",
    TEXT_HEADING    = "#1a2744",
    TEXT_H3         = "#2a3a60",
    TEXT_CODE       = "#2a3a60",
    TEXT_BQ         = "#555e7a",
    ACCENT          = "#1a2744",
    ACCENT_HOVER    = "#2a3a6a",
    ACCENT_FOCUS    = "rgba(26,39,68,0.18)",
    ACCENT_BG_HOVER = "rgba(26,39,68,0.06)",
    BTN_TEXT        = "#ffffff",
    TAB_INACTIVE    = "#7a86a8",
    ERR_BG          = "#fff0f2",
    ERR_BORDER      = "#f5b8c0",
    ERR_TEXT        = "#a01830",
    LOCK_COLOR      = "#8a94b0",
    TIER_STD_BG     = "#e8ecf8",
    TIER_STD_TEXT   = "#4a5a80",
    TIER_STD_BORDER = "#c0c8e0",
    TIER_PREM_BG    = "rgba(201,168,76,0.12)",
    TIER_PREM_TEXT  = "#7a5800",
    TIER_PREM_BORDER= "#c9a84c",
    STATS_COLOR     = "#7a86a8",
    STATS_STRONG    = "#1a2744",
    META_COLOR      = "#555e7a",
    META_STRONG     = "#7a5800",
    EMPTY_BORDER    = "#dde1ef",
    EMPTY_COLOR     = "#8a94b0",
    EMPTY_STRONG    = "#555e7a",
    EMPTY_SUB       = "#9aa4c0",
    TOGGLE_BG       = "rgba(26,39,68,0.05)",
    TOGGLE_LABEL    = "#555e7a",
)

_CSS_TEMPLATE = """
<style>
  /* ── Font catch-all — Arial everywhere except code ── */
  html, body, * {
    font-family: Arial, sans-serif !important;
  }
  pre, code, kbd, samp, .stCodeBlock * {
    font-family: 'Courier New', Courier, monospace !important;
  }
  /* ── Restore Material Symbols font for Streamlit icon elements ── */
  /* Without this the Arial catch-all turns icon ligatures into literal text  */
  /* e.g. "keyboard_arrow_right", "upload" show as plain strings              */
  [data-testid="stIconMaterial"],
  span[translate="no"] {
    font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
    font-size: 1.2rem !important;
    line-height: 1 !important;
  }

  /* ── BUG-V3: Hide sidebar collapse/expand icon text artifact ── */
  /* The "keyboard_double_arrow_right" ligature leaks as literal text when    */
  /* Material Symbols font is unavailable; hide the offending spans entirely. */
  [data-testid="collapsedControl"],
  [data-testid="collapsedControl"] span,
  [data-testid="stSidebarCollapseButton"],
  [data-testid="stSidebarCollapseButton"] span,
  button[kind="headerNoPadding"],
  button[kind="headerNoPadding"] span,
  [data-testid="baseButton-headerNoPadding"],
  [data-testid="baseButton-headerNoPadding"] span {
    font-family: 'Material Symbols Rounded', 'Material Icons', sans-serif !important;
    font-size: 1.2rem !important;
    line-height: 1 !important;
    visibility: visible !important;
  }
  /* Fallback: if font still not resolved, suppress the literal text leak */
  [data-testid="collapsedControl"] span[translate="no"],
  [data-testid="stSidebarCollapseButton"] span[translate="no"],
  button[kind="headerNoPadding"] span[translate="no"] {
    font-size: 0 !important;
    color: transparent !important;
    visibility: hidden !important;
  }
  [data-tooltip]::after { display: none !important; }

  /* ── Global ── */
  html, body,
  [data-testid="stApp"],
  [data-testid="stAppViewContainer"],
  [data-testid="stMainBlockContainer"],
  .main .block-container {
    background-color: __BG_MAIN__ !important;
    color: __TEXT_MAIN__;
  }

  /* ── Hide Streamlit chrome ── */
  #MainMenu, footer, header { visibility: hidden; }

  /* ── Page header ── */
  .pe-header {
    border-bottom: 1px solid __BORDER__;
    padding-bottom: 1.2rem;
    margin-bottom: 1.5rem;
  }
  .pe-header h1 {
    font-size: 1.65rem;
    font-weight: 700;
    color: __TEXT_HEADING__;
    letter-spacing: -0.3px;
    margin: 0 0 0.25rem 0;
  }
  .pe-header p {
    color: __TEXT_MUTED__;
    font-size: 0.88rem;
    margin: 0;
  }

  /* ── Tabs ── */
  [data-testid="stTabs"] [data-baseweb="tab-list"] {
    gap: 0;
    border-bottom: 1px solid __BORDER__;
    margin-bottom: 1.5rem;
  }
  [data-testid="stTabs"] [data-baseweb="tab"] {
    background: transparent !important;
    color: __TAB_INACTIVE__ !important;
    font-size: 0.85rem;
    font-weight: 600;
    padding: 0.6rem 1.4rem;
    border-bottom: 2px solid transparent !important;
    border-radius: 0 !important;
  }
  [data-testid="stTabs"] [aria-selected="true"] {
    color: __ACCENT__ !important;
    border-bottom: 2px solid __ACCENT__ !important;
    background: transparent !important;
  }

  /* ── Section labels ── */
  .section-label {
    font-size: 0.7rem;
    font-weight: 700;
    letter-spacing: 0.1em;
    text-transform: uppercase;
    color: __ACCENT__;
    margin-bottom: 0.4rem;
    display: block;
  }

  /* ── Form card ── */
  .form-card {
    background: __BG_PANEL__;
    border: 1px solid __BORDER__;
    border-radius: 8px;
    padding: 1.5rem 1.75rem;
    margin-bottom: 1.5rem;
  }

  /* ── Input overrides ── */
  [data-testid="stTextInput"] input,
  [data-testid="stTextArea"] textarea,
  [data-testid="stSelectbox"] select,
  [data-testid="stNumberInput"] input,
  div[data-baseweb="select"] {
    background-color: __BG_INPUT__ !important;
    border: 1px solid __BORDER_INPUT__ !important;
    border-radius: 5px !important;
    color: __TEXT_MAIN__ !important;
  }
  [data-testid="stTextInput"] input:focus,
  [data-testid="stTextArea"] textarea:focus,
  [data-testid="stNumberInput"] input:focus {
    border-color: __ACCENT__ !important;
    box-shadow: 0 0 0 2px __ACCENT_FOCUS__ !important;
  }
  label,
  .stTextInput label, .stTextArea label,
  .stSelectbox label, .stNumberInput label,
  [data-testid="stWidgetLabel"] {
    color: __TEXT_LABEL__ !important;
    font-size: 0.82rem !important;
    font-weight: 500 !important;
    text-shadow: none !important;
    -webkit-font-smoothing: antialiased !important;
  }

  /* ── Primary button — BUG-V2: force blue, prevent text doubling ── */
  [data-testid="stButton"] > button[kind="primary"],
  [data-testid="stBaseButton-primary"],
  [data-testid="stFormSubmitButton"] > button {
    background-color: #2563EB !important;
    background: #2563EB !important;
    color: #FFFFFF !important;
    border: none !important;
    border-radius: 5px !important;
    font-weight: 700 !important;
    font-size: 0.9rem !important;
    padding: 0.6rem 1.8rem !important;
    letter-spacing: 0.02em !important;
    transition: background 0.15s !important;
    text-shadow: none !important;
    -webkit-font-smoothing: antialiased !important;
  }
  [data-testid="stButton"] > button[kind="primary"]:hover,
  [data-testid="stBaseButton-primary"]:hover,
  [data-testid="stFormSubmitButton"] > button:hover {
    background-color: #1D4ED8 !important;
    background: #1D4ED8 !important;
  }

  /* ── Secondary / plain buttons ── */
  [data-testid="stButton"] > button:not([kind="primary"]) {
    color: __TEXT_LABEL__ !important;
    background: __BG_BTN_SEC__ !important;
    border: 1px solid __BORDER_INPUT__ !important;
    border-radius: 5px !important;
    text-shadow: none !important;
    -webkit-font-smoothing: antialiased !important;
  }
  [data-testid="stButton"] > button:not([kind="primary"]):hover {
    border-color: __ACCENT__ !important;
    color: __ACCENT__ !important;
  }

  /* ── Download button ── */
  [data-testid="stDownloadButton"] > button {
    background: transparent;
    border: 1px solid __ACCENT__;
    color: __ACCENT__;
    border-radius: 5px;
    font-weight: 600;
    font-size: 0.85rem;
    padding: 0.45rem 1.2rem;
  }
  [data-testid="stDownloadButton"] > button:hover {
    background: __ACCENT_BG_HOVER__;
  }

  /* ── Status / progress ── */
  [data-testid="stStatus"] {
    background: __BG_PANEL__ !important;
    border: 1px solid __BORDER__ !important;
    border-radius: 6px !important;
  }

  /* ── Brief / VCP output area ── */
  .brief-wrapper {
    background: __BG_PANEL__;
    border: 1px solid __BORDER__;
    border-radius: 8px;
    padding: 2rem 2.25rem;
  }
  .brief-wrapper h1 {
    color: __TEXT_HEADING__;
    font-size: 1.4rem;
    border-bottom: 1px solid __BORDER__;
    padding-bottom: 0.5rem;
  }
  .brief-wrapper h2 { color: __ACCENT__; font-size: 1.05rem; margin-top: 1.8rem; }
  .brief-wrapper h3 { color: __TEXT_H3__; font-size: 0.95rem; }
  .brief-wrapper p  { color: __TEXT_BODY__; line-height: 1.7; }
  .brief-wrapper li { color: __TEXT_BODY__; line-height: 1.6; }
  .brief-wrapper table { border-collapse: collapse; width: 100%; }
  .brief-wrapper th {
    background: __BG_TABLE_HDR__;
    color: __ACCENT__;
    padding: 0.5rem 0.75rem;
    font-size: 0.8rem;
    text-align: left;
  }
  .brief-wrapper td {
    padding: 0.45rem 0.75rem;
    border-bottom: 1px solid __BORDER_ROW__;
    color: __TEXT_BODY__;
    font-size: 0.85rem;
  }
  .brief-wrapper code {
    background: __BG_INPUT__;
    border-radius: 3px;
    padding: 0.1em 0.35em;
    font-size: 0.85em;
    font-family: 'Courier New', monospace !important;
  }
  .brief-wrapper blockquote {
    border-left: 3px solid #c9a84c;
    margin: 0.75rem 0;
    padding-left: 1rem;
    color: __TEXT_BQ__;
  }
  .brief-wrapper pre {
    background: __BG_INPUT__;
    border-radius: 5px;
    padding: 1rem;
    overflow-x: auto;
  }
  .brief-wrapper pre code {
    background: transparent;
    padding: 0;
    font-size: 0.82rem;
    color: __TEXT_CODE__;
    font-family: 'Courier New', monospace !important;
  }

  /* ── Divider ── */
  hr { border-color: __BORDER__ !important; }

  /* ── Sidebar divider ── */
  .sidebar-divider {
    border: 0;
    border-top: 1px solid __BORDER__ !important;
    margin: 0.75rem 0 !important;
  }

  /* ── Error box ── */
  [data-testid="stAlert"] {
    background: __ERR_BG__ !important;
    border: 1px solid __ERR_BORDER__ !important;
    color: __ERR_TEXT__ !important;
    border-radius: 6px !important;
  }

  /* ── Sidebar ── */
  [data-testid="stSidebar"],
  [data-testid="stSidebar"] > div,
  [data-testid="stSidebarContent"] {
    background-color: __BG_SIDEBAR__ !important;
    border-right: 1px solid __BORDER__ !important;
  }
  [data-testid="stSidebar"] label {
    color: __TEXT_LABEL__ !important;
    font-size: 0.82rem !important;
    text-shadow: none !important;
    -webkit-font-smoothing: antialiased !important;
  }

  /* ── Theme toggle row ── */
  .theme-toggle-row {
    display: flex;
    align-items: center;
    gap: 0.5rem;
    background: __TOGGLE_BG__;
    border: 1px solid __BORDER__;
    border-radius: 6px;
    padding: 0.4rem 0.6rem;
    margin-bottom: 0.25rem;
  }
  .theme-toggle-label {
    font-size: 0.78rem;
    color: __TOGGLE_LABEL__;
    font-weight: 600;
    letter-spacing: 0.04em;
    text-transform: uppercase;
    flex: 1;
  }

  /* ── Tier badges ── */
  .tier-badge-standard {
    display: inline-block;
    background: __TIER_STD_BG__;
    color: __TIER_STD_TEXT__;
    border: 1px solid __TIER_STD_BORDER__;
    border-radius: 4px;
    font-size: 0.7rem;
    font-weight: 700;
    padding: 0.12rem 0.5rem;
    letter-spacing: 0.06em;
    text-transform: uppercase;
  }
  .tier-badge-premium {
    display: inline-block;
    background: __TIER_PREM_BG__;
    color: __TIER_PREM_TEXT__;
    border: 1px solid __TIER_PREM_BORDER__;
    border-radius: 4px;
    font-size: 0.7rem;
    font-weight: 700;
    padding: 0.12rem 0.5rem;
    letter-spacing: 0.06em;
    text-transform: uppercase;
  }
  .lock-note {
    color: __LOCK_COLOR__;
    font-size: 0.73rem;
    font-style: italic;
    line-height: 1.2;
    margin: -0.1rem 0 0.35rem 1.6rem;
    display: block;
  }

  /* ── Sidebar checkboxes & radios — spacing ── */
  [data-testid="stSidebar"] [data-testid="stCheckbox"] {
    margin-bottom: 0.05rem;
  }
  [data-testid="stSidebar"] [data-testid="stRadio"] > div {
    gap: 0.45rem !important;
  }
  [data-testid="stSidebar"] [data-testid="stRadio"] label {
    padding: 0.25rem 0 !important;
  }

  /* ── Session stats ── */
  .session-stats {
    font-size: 0.78rem;
    color: __STATS_COLOR__;
    line-height: 1.7;
    margin: 0;
  }
  .session-stats strong { color: __STATS_STRONG__; }

  /* ── Output meta line ── */
  .output-meta {
    font-size: 0.82rem;
    color: __META_COLOR__;
    margin: 0;
  }
  .output-meta strong { color: __META_STRONG__; }

  /* ── Empty state placeholder ── */
  .empty-state {
    height: 420px;
    display: flex;
    flex-direction: column;
    align-items: center;
    justify-content: center;
    border: 1px dashed __EMPTY_BORDER__;
    border-radius: 8px;
    text-align: center;
    padding: 2rem;
    color: __EMPTY_COLOR__;
  }
  .empty-state .icon { font-size: 2.5rem; margin-bottom: 1rem; opacity: 0.45; }
  .empty-state .title {
    font-size: 0.92rem;
    font-weight: 600;
    color: __EMPTY_STRONG__;
    margin-bottom: 0.4rem;
  }
  .empty-state .sub {
    font-size: 0.8rem;
    color: __EMPTY_SUB__;
    max-width: 340px;
    line-height: 1.6;
  }
  .empty-state .sub strong { color: __EMPTY_STRONG__; }

  /* ── Captions ── */
  [data-testid="stCaptionContainer"] { color: __TEXT_MUTED__ !important; }

  /* ── Expander ── */
  [data-testid="stExpander"] {
    background: __BG_EXPANDER__ !important;
    border: 1px solid __BORDER__ !important;
    border-radius: 6px !important;
  }
  [data-testid="stExpander"] summary {
    color: __TEXT_LABEL__ !important;
    text-shadow: none !important;
    -webkit-font-smoothing: antialiased !important;
  }
</style>
"""


def _build_css(t: dict) -> str:
    css = _CSS_TEMPLATE
    for k, v in t.items():
        css = css.replace(f"__{k}__", v)
    return css


# Theme state — initialize before first CSS injection
if "sb_theme_toggle" not in st.session_state:
    st.session_state["sb_theme_toggle"] = False  # dark by default

_theme = _LIGHT if st.session_state.get("sb_theme_toggle", False) else _DARK
st.markdown(_build_css(_theme), unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Document extraction utilities
# ---------------------------------------------------------------------------

def extract_text_from_upload(uploaded_file) -> str:
    """Extract plain text from a .docx or .pdf upload."""
    if uploaded_file is None:
        return ""
    name = uploaded_file.name.lower()
    data = uploaded_file.read()
    uploaded_file.seek(0)  # reset so Streamlit can re-read if needed

    if name.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(data))
            return "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        except Exception as e:
            st.warning(f"Could not parse .docx: {e}")
            return ""

    elif name.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(data))
            pages = [page.extract_text() or "" for page in reader.pages]
            return "\n".join(pages)
        except Exception as e:
            st.warning(f"Could not parse .pdf: {e}")
            return ""

    return ""


# ---------------------------------------------------------------------------
# Word export utility
# ---------------------------------------------------------------------------

def markdown_to_docx(markdown_text: str, company_name: str, watermark: str = "") -> bytes:
    """
    Convert markdown to a styled .docx.
    Handles: # h1, ## h2, ### h3, **bold**, *italic*, tables, bullet lists, horizontal rules.
    Returns raw bytes of the .docx file.
    """
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    NAVY   = RGBColor(0x1a, 0x27, 0x4a)
    GOLD   = RGBColor(0xc9, 0xa8, 0x4c)
    BODY   = RGBColor(0x2d, 0x2d, 0x2d)
    SUBTLE = RGBColor(0x6b, 0x6b, 0x6b)

    doc = Document()

    # --- Page margins ---
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.15)
        section.right_margin  = Inches(1.15)

    # --- Base style ---
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BODY

    def set_run_formatting(run, bold=False, italic=False, color=None):
        run.font.name = "Arial"
        run.bold = bold
        run.italic = italic
        if color:
            run.font.color.rgb = color

    def add_inline_text(para, text):
        """Parse **bold**, *italic*, and plain text within a line."""
        pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|(`[^`]+`))")
        pos = 0
        for m in pattern.finditer(text):
            if m.start() > pos:
                run = para.add_run(text[pos:m.start()])
                set_run_formatting(run)
            full = m.group(0)
            if full.startswith("**"):
                run = para.add_run(m.group(2))
                set_run_formatting(run, bold=True)
            elif full.startswith("*"):
                run = para.add_run(m.group(3))
                set_run_formatting(run, italic=True)
            else:
                run = para.add_run(m.group(0).strip("`"))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            pos = m.end()
        if pos < len(text):
            run = para.add_run(text[pos:])
            set_run_formatting(run)

    def add_table_from_md(lines):
        """Parse a markdown table block and add it to the doc."""
        rows = []
        for line in lines:
            if re.match(r"^\|[-:| ]+\|$", line.strip()):
                continue  # separator row
            cells = [c.strip() for c in line.strip().strip("|").split("|")]
            rows.append(cells)
        if not rows:
            return
        cols = max(len(r) for r in rows)
        tbl = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = "Table Grid"
        for i, row_data in enumerate(rows):
            for j, cell_text in enumerate(row_data):
                if j >= cols:
                    break
                cell = tbl.cell(i, j)
                cell.text = ""
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after  = Pt(3)
                add_inline_text(para, cell_text)
                if i == 0:
                    for run in para.runs:
                        run.bold = True
                        run.font.color.rgb = NAVY
                    # shade header row
                    tc = cell._tc
                    tcPr = tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "E8ECF4")
                    shd.set(qn("w:val"), "clear")
                    tcPr.append(shd)

    # --- Parse lines ---
    lines = markdown_text.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Table detection: peek ahead to collect full table block
        if stripped.startswith("|") and stripped.endswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_table_from_md(table_lines)
            doc.add_paragraph()
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}$", stripped):
            para = doc.add_paragraph()
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:space"), "1")
            bottom.set(qn("w:color"), "C9A84C")
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # H1
        if stripped.startswith("# ") and not stripped.startswith("##"):
            para = doc.add_heading(level=1)
            para.clear()
            run = para.add_run(stripped[2:])
            run.font.name = "Arial"
            run.font.size = Pt(18)
            run.font.color.rgb = NAVY
            run.bold = True
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(8)
            i += 1
            continue

        # H2
        if stripped.startswith("## "):
            para = doc.add_heading(level=2)
            para.clear()
            run = para.add_run(stripped[3:])
            run.font.name = "Arial"
            run.font.size = Pt(13)
            run.font.color.rgb = GOLD
            run.bold = True
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # H3
        if stripped.startswith("### "):
            para = doc.add_heading(level=3)
            para.clear()
            run = para.add_run(stripped[4:])
            run.font.name = "Arial"
            run.font.size = Pt(11)
            run.font.color.rgb = NAVY
            run.bold = True
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after  = Pt(2)
            i += 1
            continue

        # Blockquote
        if stripped.startswith("> "):
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.4)
            run = para.add_run(stripped[2:])
            run.font.color.rgb = SUBTLE
            run.italic = True
            para.paragraph_format.space_after = Pt(6)
            i += 1
            continue

        # Bullet list (- or *)
        if re.match(r"^[-*]\s+", stripped):
            para = doc.add_paragraph(style="List Bullet")
            add_inline_text(para, re.sub(r"^[-*]\s+", "", stripped))
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Numbered list
        if re.match(r"^\d+\.\s+", stripped):
            para = doc.add_paragraph(style="List Number")
            add_inline_text(para, re.sub(r"^\d+\.\s+", "", stripped))
            para.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Code block
        if stripped.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # skip closing ```
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.3)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run("\n".join(code_lines))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # Empty line
        if not stripped:
            i += 1
            continue

        # Normal paragraph
        para = doc.add_paragraph()
        add_inline_text(para, stripped)
        para.paragraph_format.space_after = Pt(5)
        i += 1

    # Watermark footer for standard tier
    if watermark:
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(24)
        para.paragraph_format.space_after  = Pt(0)
        from docx.oxml.ns import qn as _qn
        from docx.oxml import OxmlElement as _OE
        pPr = para._p.get_or_add_pPr()
        pBdr = _OE("w:pBdr")
        top = _OE("w:top")
        top.set(_qn("w:val"), "single")
        top.set(_qn("w:sz"), "4")
        top.set(_qn("w:space"), "1")
        top.set(_qn("w:color"), "AAAAAA")
        pBdr.append(top)
        pPr.append(pBdr)
        run = para.add_run(watermark)
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.italic = True
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PPT export helpers
# ---------------------------------------------------------------------------

def _md_to_pptx_sections(md_text: str) -> dict[str, str]:
    """
    Parse a numbered-section markdown brief into an ordered {title: content} dict
    suitable for passing to export_pptx() as the content argument.
    Skips metadata sections (Data Sources, Generated by).
    """
    pattern = re.compile(r"^##\s+(?:\d+\.\s+)?(.+?)$", re.MULTILINE)
    matches = list(pattern.finditer(md_text))
    sections: dict[str, str] = {}
    for i, m in enumerate(matches):
        title = m.group(1).strip()
        if re.search(r"data sources|generated by", title, re.IGNORECASE):
            continue
        start   = m.end()
        end     = matches[i + 1].start() if i + 1 < len(matches) else len(md_text)
        content = md_text[start:end].strip()
        if content and len(content) > 30:
            sections[title] = content
    return sections


def _render_pptx_export_options(tab_prefix: str) -> None:
    """
    Collapsible Export Options section for PPT template upload.
    Reads/writes st.session_state["pptx_template"] and ["pptx_template_name"].
    Call inside an expander in each tab's form column.
    """
    st.caption(
        "Upload your firm's PowerPoint template to match its formatting. "
        "Leave blank to use default styling."
    )
    pptx_file = st.file_uploader(
        "Upload PPT template (optional)",
        type=["pptx"],
        key=f"pptx_tpl_{tab_prefix}",
        label_visibility="collapsed",
    )
    if pptx_file is not None:
        st.session_state["pptx_template"]      = pptx_file.read()
        st.session_state["pptx_template_name"] = pptx_file.name

    tpl_name = st.session_state.get("pptx_template_name")
    if tpl_name:
        st.markdown(
            f'<span style="color:#2e7d32;font-size:0.82rem;font-weight:600;">'
            f'✓ Template loaded: {tpl_name}</span>',
            unsafe_allow_html=True,
        )
    else:
        st.caption("Using default PE Ops styling (Arial, navy/white)")


# ---------------------------------------------------------------------------
# Tier & sidebar state initialization
# ---------------------------------------------------------------------------
sys.path.insert(0, str(BASE_DIR / "src"))
from tier import (  # noqa: E402
    check_premium_password,
    is_source_locked,
    STANDARD_BRIEF_LIMIT,
    WATERMARK_TEXT,
)

if "tier" not in st.session_state:
    st.session_state["tier"] = "standard"
if "brief_count" not in st.session_state:
    st.session_state["brief_count"] = 0


# ---------------------------------------------------------------------------
# Sidebar — Theme toggle · Tier gate · Model Quality · Data Sources
# ---------------------------------------------------------------------------

with st.sidebar:
    tier = st.session_state["tier"]

    # ── Theme toggle (top of sidebar) ────────────────────────────────────────
    toggle_label = "Light mode" if not st.session_state.get("sb_theme_toggle") else "Dark mode"
    st.markdown(
        f'<div class="theme-toggle-row">'
        f'<span class="theme-toggle-label">{toggle_label}</span>',
        unsafe_allow_html=True,
    )
    st.toggle(
        "Switch theme",
        key="sb_theme_toggle",
        label_visibility="collapsed",
    )
    st.markdown('</div>', unsafe_allow_html=True)

    st.markdown('<hr class="sidebar-divider">', unsafe_allow_html=True)

    # ── Pricing tier ─────────────────────────────────────────────────────────
    st.markdown('<span class="section-label">Pricing Tier</span>', unsafe_allow_html=True)

    if tier == "premium":
        st.markdown(
            '<span class="tier-badge-premium">Premium</span>',
            unsafe_allow_html=True,
        )
        st.caption("Opus unlocked · Unlimited briefs · All data sources")
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Downgrade to Standard", key="sb_downgrade", use_container_width=True):
            st.session_state["tier"] = "standard"
            st.session_state["brief_count"] = 0
            st.rerun()
    else:
        briefs_used = st.session_state["brief_count"]
        briefs_left = max(0, STANDARD_BRIEF_LIMIT - briefs_used)
        st.markdown(
            f'<span class="tier-badge-standard">Standard</span>',
            unsafe_allow_html=True,
        )
        st.caption(f"{briefs_left} of {STANDARD_BRIEF_LIMIT} briefs remaining this session")
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("Unlock Premium", expanded=False):
            pw_input = st.text_input(
                "Password",
                type="password",
                placeholder="Enter premium password",
                key="sb_pw_input",
            )
            if st.button("Unlock", key="sb_unlock_btn", use_container_width=True):
                if check_premium_password(pw_input):
                    st.session_state["tier"] = "premium"
                    st.session_state["brief_count"] = 0
                    st.rerun()
                else:
                    st.error("Incorrect password.")

    st.markdown('<hr class="sidebar-divider">', unsafe_allow_html=True)

    # ── Model quality ─────────────────────────────────────────────────────────
    st.markdown('<span class="section-label">Model Quality</span>', unsafe_allow_html=True)

    if tier == "standard":
        st.caption("Sonnet only on Standard tier")
        st.session_state["sb_model_mode"] = "standard"
        st.radio(
            "Model",
            options=["Sonnet — Standard (~$0.20/brief)"],
            index=0,
            key="sb_model_radio_std",
            label_visibility="collapsed",
        )
    else:
        model_selection = st.radio(
            "Model",
            options=[
                "Sonnet — Standard (~$0.20/brief)",
                "Opus — Premium (~$3–5/brief)",
            ],
            index=0,
            key="sb_model_radio_prem",
            label_visibility="collapsed",
        )
        st.session_state["sb_model_mode"] = (
            "premium" if "Opus" in model_selection else "standard"
        )

    model_mode = st.session_state.get("sb_model_mode", "standard")

    st.markdown('<hr class="sidebar-divider">', unsafe_allow_html=True)

    # ── Data sources ──────────────────────────────────────────────────────────
    st.markdown('<span class="section-label">Data Sources</span>', unsafe_allow_html=True)

    _ds_defs = [
        ("sec_edgar",     "SEC EDGAR",                    False),
        ("yahoo_finance", "Yahoo Finance",                False),
        ("bls",           "BLS Labor Data",               False),
        ("news",          "News Sweep (accredited)",      False),
        ("damodaran",     "Damodaran Multiples",          True),
        ("naver_finance", "Naver Finance (Korea/Asia)",   True),
    ]

    selected_sources: list[str] = []
    for src_key, src_label, premium_only in _ds_defs:
        locked = premium_only and tier == "standard"
        label = f"🔒 {src_label}" if locked else src_label
        help_text = "Requires Premium tier" if locked else None
        cb_val = st.checkbox(
            label,
            value=(not locked),
            disabled=locked,
            key=f"ds_{src_key}",
            help=help_text,
        )
        if cb_val and not locked:
            selected_sources.append(src_key)

    st.session_state["_selected_sources"] = selected_sources

    st.markdown('<hr class="sidebar-divider">', unsafe_allow_html=True)

    # ── Session stats ─────────────────────────────────────────────────────────
    st.markdown('<span class="section-label">Session</span>', unsafe_allow_html=True)
    bc = st.session_state["brief_count"]
    tier_label = "Premium" if tier == "premium" else "Standard"
    st.markdown(
        f'<p class="session-stats">'
        f'Briefs generated: <strong>{bc}</strong><br>'
        f'Tier: <strong>{tier_label}</strong>'
        f'</p>',
        unsafe_allow_html=True,
    )


# ---------------------------------------------------------------------------
# Header
# ---------------------------------------------------------------------------
st.markdown("""
<div class="pe-header">
  <h1>PE Ops Tool Suite</h1>
  <p>Middle market buyout · Operational intelligence and value creation planning</p>
</div>
""", unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Utility: run async coroutine in a worker thread
# ---------------------------------------------------------------------------
def run_async_in_thread(coro, result_queue: queue.Queue):
    loop = asyncio.new_event_loop()
    asyncio.set_event_loop(loop)
    try:
        result = loop.run_until_complete(coro)
        result_queue.put(("ok", result))
    except Exception as exc:
        result_queue.put(("err", exc))
    finally:
        loop.close()


# ---------------------------------------------------------------------------
# Shared: output display block (used by both tabs)
# ---------------------------------------------------------------------------
def render_output_block(result_key: str, inputs_key: str, error_key: str,
                        generating_key: str, label: str, file_suffix: str,
                        include_lbo: bool = False, docx_watermark: str = "",
                        pptx_tab: str = ""):
    """Render the result/error/empty state for a generation module.

    include_lbo: when True and a result exists, generate and surface an LBO
                 model Excel download alongside the standard Markdown / Word buttons.
    pptx_tab:    when non-empty ("diligence" or "vcp"), adds a PPT export button.
    """

    if st.session_state.get(error_key):
        st.error(f"**Error:** {st.session_state[error_key]}")
        if st.button("Try again", key=f"{result_key}_retry"):
            st.session_state[error_key] = None
            st.rerun()

    elif st.session_state.get(result_key):
        output_text = st.session_state[result_key]
        inputs = st.session_state.get(inputs_key, {})
        company = inputs.get("company_name", "Company")
        today = datetime.today().strftime("%Y-%m-%d")
        slug = company.lower().replace(" ", "_")
        md_filename   = f"{slug}_{file_suffix}_{today}.md"
        docx_filename = f"{slug}_{file_suffix}_{today}.docx"

        # Optionally generate LBO Excel bytes (fast — no API call)
        lbo_bytes = None
        if include_lbo:
            try:
                from lbo_model import build_lbo_excel  # noqa: E402
                lbo_bytes = build_lbo_excel(
                    company_name=company,
                    industry=inputs.get("industry", ""),
                    ev_range=inputs.get("ev_range", "$100–250M"),
                )
            except Exception as exc:
                st.warning(f"LBO model could not be generated: {exc}")

        # Build download column layout
        n_extra = (1 if lbo_bytes else 0) + (1 if pptx_tab else 0)
        if n_extra == 2:
            meta_col, dl_md_col, dl_docx_col, dl_lbo_col, dl_pptx_col = st.columns([2, 1, 1, 1, 1])
        elif n_extra == 1 and lbo_bytes:
            meta_col, dl_md_col, dl_docx_col, dl_lbo_col = st.columns([2, 1, 1, 1])
            dl_pptx_col = None
        elif n_extra == 1 and pptx_tab:
            meta_col, dl_md_col, dl_docx_col, dl_pptx_col = st.columns([2, 1, 1, 1])
            dl_lbo_col = None
        else:
            meta_col, dl_md_col, dl_docx_col = st.columns([2, 1, 1])
            dl_lbo_col = dl_pptx_col = None

        with meta_col:
            st.markdown(
                f'<p class="output-meta">'
                f'<strong>{company}</strong> · '
                f'{inputs.get("industry", "")} · {today}'
                f'</p>',
                unsafe_allow_html=True,
            )
        with dl_md_col:
            st.download_button(
                label="Download Markdown",
                data=output_text.encode("utf-8"),
                file_name=md_filename,
                mime="text/markdown",
                use_container_width=True,
                key=f"{result_key}_dl_md",
            )
        with dl_docx_col:
            docx_bytes = markdown_to_docx(output_text, company, watermark=docx_watermark)
            st.download_button(
                label="Download as Word",
                data=docx_bytes,
                file_name=docx_filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key=f"{result_key}_dl_docx",
            )
        if dl_lbo_col and lbo_bytes:
            with dl_lbo_col:
                st.download_button(
                    label="Download LBO Model",
                    data=lbo_bytes,
                    file_name=f"{slug}_lbo_model_{today}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True,
                    key=f"{result_key}_dl_lbo",
                )
        if dl_pptx_col and pptx_tab:
            with dl_pptx_col:
                try:
                    from pptx_export import export_pptx  # noqa: E402
                    _tpl_bytes = st.session_state.get("pptx_template")
                    _tpl_io    = io.BytesIO(_tpl_bytes) if _tpl_bytes else None
                    _pptx_secs = _md_to_pptx_sections(output_text)
                    _pptx_buf  = export_pptx(_pptx_secs, pptx_tab, company, template=_tpl_io)
                    st.download_button(
                        label="Export as PowerPoint",
                        data=_pptx_buf.getvalue(),
                        file_name=f"{slug}_{file_suffix}_{today}.pptx",
                        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                        use_container_width=True,
                        key=f"{result_key}_dl_pptx",
                    )
                except Exception as _pptx_err:
                    st.caption(f"PPT unavailable: {_pptx_err}")

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="brief-wrapper">', unsafe_allow_html=True)
        st.markdown(output_text, unsafe_allow_html=False)
        st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        if st.button(f"Generate another {label}", use_container_width=False, key=f"{result_key}_reset"):
            for key in [result_key, error_key, f"{result_key}_path", inputs_key, generating_key]:
                st.session_state.pop(key, None)
            st.rerun()

    elif not st.session_state.get(generating_key) and not st.session_state.get(error_key):
        icon = "📋" if label == "brief" else "📊"
        action = "Generate Brief" if label == "brief" else "Generate VCP"
        st.markdown(f"""
<div class="empty-state">
  <div class="icon">{icon}</div>
  <div class="title">Output will appear here</div>
  <div class="sub">Fill in the details on the left and click <strong>{action}</strong>.</div>
</div>
""", unsafe_allow_html=True)


# ===========================================================================
# TAB LAYOUT
# ===========================================================================
tab1, tab2, tab3 = st.tabs(["Deal Diligence Brief", "Value Creation Planner", "100-Day Plan"])


# ===========================================================================
# TAB 1 — Deal Diligence Brief
# ===========================================================================
with tab1:
    col_form, col_out = st.columns([1, 1.6], gap="large")

    with col_form:
        st.markdown('<div class="section-label">Target Details</div>', unsafe_allow_html=True)

        b_company = st.text_input(
            "Company name *",
            placeholder="e.g. Acme Packaging",
            key="b_company_name",
        )
        b_description = st.text_area(
            "Company description *",
            placeholder=(
                "2–5 sentences on the business: what they do, who they serve, "
                "revenue model, and any known ownership history."
            ),
            height=130,
            key="b_description",
        )
        b_industry = st.text_input(
            "Industry vertical *",
            placeholder="e.g. Industrial Packaging, Healthcare Services, Business Services",
            key="b_industry",
        )
        b_ev_range = st.text_input(
            "EV range *",
            placeholder="e.g. $150M–$200M or ~$175M",
            value="$100–250M",
            key="b_ev_range",
        )
        b_deal_type = st.selectbox(
            "Deal type",
            options=[
                "Standard (Mfg / Services / Retail)",
                "Financial Services / Fintech",
                "Real Estate",
                "Technology / SaaS",
            ],
            index=0,
            key="b_deal_type",
            help=(
                "Sets which deal-type specific modules appear in the checklist below. "
                "Financial Services adds Credit Metrics; Technology / SaaS adds SaaS Metrics."
            ),
        )
        b_brief_style = st.radio(
            "Brief style",
            options=["Long Form (~10,000 words)", "Bullet (~5,000 words)"],
            index=0,
            key="b_brief_style",
            horizontal=True,
            help=(
                "Long Form: full prose sections with detailed analysis. "
                "Bullet: concise bullets-only format, ~5,000 words max."
            ),
        )

        b_notes = st.text_area(
            "Context notes (optional)",
            placeholder=(
                "Deal-specific context: ownership history, known ops issues, "
                "thesis hooks, key-man concerns, customer concentration flags…"
            ),
            height=100,
            key="b_context_notes",
        )

        # ── Module selector ──────────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Sections to Generate</div>', unsafe_allow_html=True)

        MODULE_OPTIONS = [
            ("exec_summary",        "Executive Summary"),
            ("risk_flags",          "Operational Risk Flags"),
            ("it_systems",          "IT & Systems Maturity"),
            ("comps_benchmarks",    "Comparable Companies & Benchmarks"),
            ("value_creation",      "Value Creation Opportunities"),
            ("100_day_plan",        "100-Day Plan"),
            ("diligence_questions", "Diligence Questions"),
        ]

        selected_modules = []
        cb_cols = st.columns(2)
        for idx, (key, label_text) in enumerate(MODULE_OPTIONS):
            col = cb_cols[idx % 2]
            checked = col.checkbox(label_text, value=True, key=f"b_mod_{key}")
            if checked:
                selected_modules.append(key)

        # ── Optional modules (off by default) ────────────────────────────────
        st.markdown(
            '<div style="border-top:1px solid #2a2d35;margin:0.75rem 0 0.5rem 0;"></div>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="section-label">Optional Modules</div>', unsafe_allow_html=True)

        _opt_modules = [
            ("functional_scorecards",    "Functional Scorecards (Ops / IT / Commercial / Talent)"),
            ("transaction_structure",    "Transaction Structure Analysis"),
            ("change_my_view",           "What Would Change My View"),
            ("investment_recommendation","Investment Recommendation"),
        ]

        # Synergy Analysis only when context notes suggest M&A / add-on scenario
        _synergy_keywords = [
            "acquisition", "acquire", "merger", "add-on", "addon", "add on",
            "bolt-on", "bolton", "roll-up", "rollup",
        ]
        _show_synergy = bool(b_notes) and any(
            kw in b_notes.lower() for kw in _synergy_keywords
        )
        if _show_synergy:
            _opt_modules.insert(2, ("synergy_analysis", "Synergy Analysis (M&A / add-on)"))

        for key, label_text in _opt_modules:
            if st.checkbox(label_text, value=False, key=f"b_mod_{key}"):
                selected_modules.append(key)

        # Deal-type specific modules
        if "Financial Services" in b_deal_type:
            if st.checkbox(
                "Credit & Financial Metrics",
                value=False,
                key="b_mod_credit_metrics",
                help="Loss curves, NIM, funding stack, regulatory trajectory, vintage analysis.",
            ):
                selected_modules.append("credit_metrics")
        elif "Technology" in b_deal_type or "SaaS" in b_deal_type:
            if st.checkbox(
                "SaaS & Technology Metrics",
                value=False,
                key="b_mod_saas_metrics",
                help="ARR, NRR, CAC/LTV, churn analysis, technology moat assessment.",
            ):
                selected_modules.append("saas_metrics")
        elif "Real Estate" in b_deal_type:
            st.caption("Real Estate modules — coming soon")

        # ── Add-ons ──────────────────────────────────────────────────────────
        st.markdown(
            '<div style="border-top:1px solid #2a2d35;margin:0.75rem 0 0.5rem 0;"></div>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="section-label">Add-Ons</div>', unsafe_allow_html=True)
        b_include_lbo = st.checkbox(
            "LBO Model Starter (Excel)",
            value=False,
            key="b_mod_lbo_model",
            help=(
                "Generates a pre-populated 4-sheet Excel workbook "
                "(Assumptions, Income Statement, Debt Schedule, Returns Summary) "
                "seeded from the brief inputs. All assumptions are placeholders."
            ),
        )

        # ── Document upload ──────────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Style Reference (optional)</div>', unsafe_allow_html=True)
        b_upload = st.file_uploader(
            "Upload a .docx or .pdf to mirror its structure and tone",
            type=["docx", "pdf"],
            key="b_upload",
            label_visibility="collapsed",
        )
        if b_upload:
            st.caption(f"Uploaded: {b_upload.name}")

        # ── Export Options ────────────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("Export Options", expanded=False):
            _render_pptx_export_options("b")

        st.markdown("<br>", unsafe_allow_html=True)
        b_generate = st.button(
            "Generate Brief",
            type="primary",
            use_container_width=True,
            disabled=st.session_state.get("b_generating", False),
            key="b_generate_btn",
        )

    with col_out:
        # Validation & trigger
        if b_generate:
            # Standard tier brief limit check
            _cur_tier = st.session_state.get("tier", "standard")
            _brief_cnt = st.session_state.get("brief_count", 0)
            if _cur_tier == "standard" and _brief_cnt >= STANDARD_BRIEF_LIMIT:
                st.error(
                    f"Standard tier limit reached ({STANDARD_BRIEF_LIMIT} briefs/session). "
                    "Unlock Premium in the sidebar for unlimited briefs."
                )
            else:
                missing = []
                if not b_company.strip():
                    missing.append("Company name")
                if not b_description.strip():
                    missing.append("Company description")
                if not b_industry.strip():
                    missing.append("Industry vertical")
                if not selected_modules:
                    missing.append("at least one section")

                if missing:
                    st.error(f"Please fill in / select: {', '.join(missing)}")
                else:
                    style_ref = extract_text_from_upload(b_upload) if b_upload else ""
                    st.session_state["b_generating"] = True
                    st.session_state["b_result"] = None
                    st.session_state["b_error"] = None
                    st.session_state["b_inputs"] = {
                        "company_name": b_company.strip(),
                        "description": b_description.strip(),
                        "industry": b_industry.strip(),
                        "ev_range": b_ev_range,
                        "deal_type": b_deal_type,
                        "context_notes": b_notes.strip(),
                        "modules": selected_modules,
                        "style_reference": style_ref,
                        "include_lbo": b_include_lbo,
                        "data_sources": st.session_state.get(
                            "_selected_sources", ["sec_edgar", "yahoo_finance", "bls", "news"]
                        ),
                        "model_mode": st.session_state.get("sb_model_mode", "standard"),
                        "brief_style": "bullet" if "Bullet" in b_brief_style else "long_form",
                    }
                    st.rerun()

        # Active generation
        if st.session_state.get("b_generating") and not st.session_state.get("b_result"):
            inputs = st.session_state["b_inputs"]
            from generate_brief import generate_brief  # noqa: E402

            _active_mods = inputs.get("modules", [])
            steps = ["Web research and data pulls (SEC EDGAR, BLS, news)…"]
            if "exec_summary"        in _active_mods: steps.append("Generating executive summary…")
            if "risk_flags"          in _active_mods: steps.append("Assessing operational risk flags…")
            if "comps_benchmarks"    in _active_mods: steps.append("Building comps and benchmarks…")
            if "it_systems"          in _active_mods: steps.append("Evaluating IT systems maturity…")
            if "value_creation"      in _active_mods: steps.append("Mapping value creation levers…")
            if "100_day_plan"        in _active_mods: steps.append("Drafting 100-day plan…")
            if "diligence_questions" in _active_mods: steps.append("Compiling diligence questions…")
            if "transaction_structure"    in _active_mods: steps.append("Analyzing transaction structure…")
            if "synergy_analysis"         in _active_mods: steps.append("Running synergy analysis…")
            if "change_my_view"           in _active_mods: steps.append("Writing What Would Change My View…")
            if "investment_recommendation" in _active_mods: steps.append("Drafting investment recommendation…")
            if "credit_metrics"           in _active_mods: steps.append("Assessing credit & financial metrics…")
            if "saas_metrics"             in _active_mods: steps.append("Evaluating SaaS & technology metrics…")
            if "functional_scorecards"    in _active_mods: steps.append("Building functional scorecards…")
            steps.append("Assembling final brief…")

            with st.status("Generating brief — this takes 1–2 minutes…", expanded=True) as status_box:
                for s in steps:
                    status_box.write(f"⏳ {s}")

                result_q: queue.Queue = queue.Queue()
                t = threading.Thread(
                    target=run_async_in_thread,
                    args=(
                        generate_brief(
                            company_name=inputs["company_name"],
                            description=inputs["description"],
                            industry=inputs["industry"],
                            ev_range=inputs["ev_range"],
                            context_notes=inputs["context_notes"],
                            modules=inputs.get("modules"),
                            style_reference=inputs.get("style_reference", ""),
                            data_sources=inputs.get("data_sources"),
                            model_mode=inputs.get("model_mode", "standard"),
                            deal_type=inputs.get("deal_type", "Standard (Mfg / Services / Retail)"),
                            brief_style=inputs.get("brief_style", "long_form"),
                        ),
                        result_q,
                    ),
                    daemon=True,
                )
                t.start()
                t.join()

                kind, payload = result_q.get_nowait()

                if kind == "ok":
                    brief_text = Path(payload).read_text(encoding="utf-8")
                    st.session_state["b_result"] = brief_text
                    st.session_state["b_result_path"] = payload
                    st.session_state["b_generating"] = False
                    st.session_state["brief_count"] = st.session_state.get("brief_count", 0) + 1
                    # Extract Value Creation section → pre-fill Tab 3 levers field
                    _vco_match = re.search(
                        r"##\s+\d+\.\s+Value Creation Opportunities\n+(.*?)(?=\n##\s+\d+\.|\Z)",
                        brief_text,
                        re.DOTALL,
                    )
                    if _vco_match:
                        _vco_raw = _vco_match.group(1).strip()
                        # Extract bullet lines as discrete levers
                        _levers = []
                        for _ln in _vco_raw.splitlines():
                            _s = _ln.strip()
                            if _s.startswith(("- ", "* ", "• ")):
                                _content = _s[2:].strip()
                                if 10 < len(_content) < 200:
                                    _levers.append(_content)
                        st.session_state["vco_for_100day"] = "\n".join(_levers[:12]) or _vco_raw[:800]
                    st.session_state["tab1_run_at"] = datetime.now().isoformat()
                    status_box.update(label="Brief ready.", state="complete", expanded=False)
                    st.rerun()
                else:
                    st.session_state["b_error"] = str(payload)
                    st.session_state["b_generating"] = False
                    status_box.update(label="Generation failed.", state="error", expanded=False)
                    st.rerun()

        _b_watermark = (
            WATERMARK_TEXT
            if st.session_state.get("tier", "standard") == "standard"
            else ""
        )
        render_output_block(
            result_key="b_result",
            inputs_key="b_inputs",
            error_key="b_error",
            generating_key="b_generating",
            label="brief",
            file_suffix="ops_brief",
            include_lbo=st.session_state.get("b_inputs", {}).get("include_lbo", False),
            docx_watermark=_b_watermark,
            pptx_tab="diligence",
        )


# ===========================================================================
# TAB 2 — Value Creation Planner
# ===========================================================================
with tab2:
    col_vcp_form, col_vcp_out = st.columns([1, 1.6], gap="large")

    with col_vcp_form:
        st.markdown('<div class="section-label">Company & Deal Inputs</div>', unsafe_allow_html=True)

        v_company = st.text_input(
            "Company name *",
            placeholder="e.g. Acme Packaging",
            key="v_company_name",
        )
        v_description = st.text_area(
            "Company description *",
            placeholder="What the business does, who it serves, revenue model.",
            height=100,
            key="v_description",
        )
        v_industry = st.text_input(
            "Industry vertical *",
            placeholder="e.g. Industrial Packaging, Healthcare Services",
            key="v_industry",
        )

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Financials</div>', unsafe_allow_html=True)

        fin_col1, fin_col2 = st.columns(2)
        with fin_col1:
            v_revenue = st.number_input(
                "Current Revenue ($M) *",
                min_value=0.0,
                value=None,
                placeholder="e.g. 85.0",
                step=1.0,
                format="%.1f",
                key="v_revenue",
            )
        with fin_col2:
            v_ebitda = st.number_input(
                "Current EBITDA ($M) *",
                min_value=0.0,
                value=None,
                placeholder="e.g. 14.0",
                step=0.5,
                format="%.1f",
                key="v_ebitda",
            )

        hold_col, target_col = st.columns(2)
        with hold_col:
            v_hold = st.selectbox(
                "Target hold period *",
                options=[3, 4, 5],
                index=2,
                key="v_hold_period",
            )
        with target_col:
            v_target_ebitda = st.number_input(
                "Target EBITDA at exit ($M) *",
                min_value=0.0,
                value=None,
                placeholder="e.g. 28.0",
                step=0.5,
                format="%.1f",
                key="v_target_ebitda",
            )

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Thesis & Context</div>', unsafe_allow_html=True)

        v_thesis = st.text_area(
            "PE sponsor's investment thesis *",
            placeholder=(
                "What's the core value creation hypothesis? "
                "e.g. Fragmented market with pricing power, undermanaged working capital, "
                "greenfield expansion opportunity into adjacent geographies."
            ),
            height=110,
            key="v_thesis",
        )
        v_challenges = st.text_area(
            "Known operational challenges",
            placeholder=(
                "Issues surfaced in diligence or known pre-close: "
                "e.g. No CFO in seat, ERP is 15 years old, top customer is 38% of revenue, "
                "margins compressed 4pp over 3 years."
            ),
            height=110,
            key="v_challenges",
        )

        # ── Optional modules ─────────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown(
            '<div style="border-top:1px solid #2a2d35;margin:0.75rem 0 0.5rem 0;"></div>',
            unsafe_allow_html=True,
        )
        st.markdown('<div class="section-label">Optional Modules</div>', unsafe_allow_html=True)
        v_include_scorecards = st.checkbox(
            "Functional Scorecards (Ops / IT / Commercial / Talent)",
            value=False,
            key="v_mod_functional_scorecards",
            help=(
                "Generates four R/Y/G maturity scorecards (15 dimensions each) covering "
                "Operations, IT & Systems, Commercial, and Talent/HR. "
                "Adds ~30–60 seconds to generation time."
            ),
        )

        # ── Document upload ──────────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Style Reference (optional)</div>', unsafe_allow_html=True)
        v_upload = st.file_uploader(
            "Upload a .docx or .pdf to mirror its structure and tone",
            type=["docx", "pdf"],
            key="v_upload",
            label_visibility="collapsed",
        )
        if v_upload:
            st.caption(f"Uploaded: {v_upload.name}")

        # ── Export Options ────────────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("Export Options", expanded=False):
            _render_pptx_export_options("v")

        st.markdown("<br>", unsafe_allow_html=True)
        v_generate = st.button(
            "Generate VCP",
            type="primary",
            use_container_width=True,
            disabled=st.session_state.get("v_generating", False),
            key="v_generate_btn",
        )

    with col_vcp_out:
        # Validation & trigger
        if v_generate:
            missing = []
            if not v_company.strip():
                missing.append("Company name")
            if not v_description.strip():
                missing.append("Company description")
            if not v_industry.strip():
                missing.append("Industry vertical")
            if v_revenue is None or v_revenue <= 0:
                missing.append("Current Revenue")
            if v_ebitda is None or v_ebitda <= 0:
                missing.append("Current EBITDA")
            if v_target_ebitda is None or v_target_ebitda <= 0:
                missing.append("Target EBITDA at exit")
            if not v_thesis.strip():
                missing.append("Investment thesis")

            if missing:
                st.error(f"Please fill in: {', '.join(missing)}")
            else:
                style_ref = extract_text_from_upload(v_upload) if v_upload else ""
                st.session_state["v_generating"] = True
                st.session_state["v_result"] = None
                st.session_state["v_error"] = None
                st.session_state["v_inputs"] = {
                    "company_name": v_company.strip(),
                    "description": v_description.strip(),
                    "industry": v_industry.strip(),
                    "current_revenue": float(v_revenue),
                    "current_ebitda": float(v_ebitda),
                    "investment_thesis": v_thesis.strip(),
                    "operational_challenges": v_challenges.strip() or "None specified.",
                    "hold_period": int(v_hold),
                    "target_ebitda": float(v_target_ebitda),
                    "style_reference": style_ref,
                    "include_scorecards": v_include_scorecards,
                }
                st.rerun()

        # Active VCP generation
        if st.session_state.get("v_generating") and not st.session_state.get("v_result"):
            inputs = st.session_state["v_inputs"]
            from generate_vcp import generate_vcp  # noqa: E402

            ebitda_uplift = inputs["target_ebitda"] - inputs["current_ebitda"]
            steps = [
                f"Modeling EBITDA bridge: +${ebitda_uplift:.1f}M over {inputs['hold_period']} years…",
                "Building top 5 value creation workstreams…",
                "Generating quick win tracker (90-day initiatives)…",
                "Designing KPI dashboard…",
                "Writing 100-day sprint plan…",
                "Identifying organizational capability gaps…",
            ]
            if inputs.get("include_scorecards"):
                steps.append("Building functional scorecards (Ops / IT / Commercial / Talent)…")
            steps.append("Assembling final VCP…")

            with st.status("Generating VCP — typically 30–60 seconds…", expanded=True) as status_box:
                for s in steps:
                    status_box.write(f"⏳ {s}")

                result_q: queue.Queue = queue.Queue()
                t = threading.Thread(
                    target=run_async_in_thread,
                    args=(
                        generate_vcp(
                            company_name=inputs["company_name"],
                            description=inputs["description"],
                            industry=inputs["industry"],
                            current_revenue=inputs["current_revenue"],
                            current_ebitda=inputs["current_ebitda"],
                            investment_thesis=inputs["investment_thesis"],
                            operational_challenges=inputs["operational_challenges"],
                            hold_period=inputs["hold_period"],
                            target_ebitda=inputs["target_ebitda"],
                            style_reference=inputs.get("style_reference", ""),
                            include_scorecards=inputs.get("include_scorecards", False),
                        ),
                        result_q,
                    ),
                    daemon=True,
                )
                t.start()
                t.join()

                kind, payload = result_q.get_nowait()

                if kind == "ok":
                    vcp_text = Path(payload).read_text(encoding="utf-8")
                    st.session_state["v_result"] = vcp_text
                    st.session_state["v_result_path"] = payload
                    st.session_state["v_generating"] = False
                    status_box.update(label="VCP ready.", state="complete", expanded=False)
                    st.rerun()
                else:
                    st.session_state["v_error"] = str(payload)
                    st.session_state["v_generating"] = False
                    status_box.update(label="Generation failed.", state="error", expanded=False)
                    st.rerun()

        render_output_block(
            result_key="v_result",
            inputs_key="v_inputs",
            error_key="v_error",
            generating_key="v_generating",
            label="VCP",
            file_suffix="vcp",
            pptx_tab="vcp",
        )


# ===========================================================================
# TAB 3 — 100-Day Plan
# ===========================================================================

_PLAN_SECTION_TITLES = {
    "workstreams":       "Workstreams",
    "resource_plan":     "Resource Plan",
    "csuite_assessment": "C-Suite Assessment",
    "org_chart":         "Org Chart (18-Month Target)",
    "ebitda_bridge":     "EBITDA Bridge",
}


def _render_100day_output() -> None:
    if st.session_state.get("p_error"):
        st.error(f"**Error:** {st.session_state['p_error']}")
        if st.button("Try again", key="p_retry"):
            st.session_state["p_error"] = None
            st.rerun()

    elif st.session_state.get("p_result"):
        sections = st.session_state["p_result"]
        inputs   = st.session_state.get("p_inputs", {})
        company  = inputs.get("company_name", "Company")
        today    = datetime.today().strftime("%Y-%m-%d")
        slug     = company.lower().replace(" ", "_")

        # Assemble combined markdown for file exports
        md_lines = [f"# 100-Day Operational Plan: {company}", "", f"**Generated:** {today}", ""]
        for key, title in _PLAN_SECTION_TITLES.items():
            content = sections.get(key, "")
            if content and not content.startswith("_Error"):
                md_lines += [f"## {title}", "", content, "", "---", ""]
        full_md = "\n".join(md_lines)

        # Export row
        meta_col, dl_md_col, dl_docx_col, dl_pptx_col = st.columns([2, 1, 1, 1])
        with meta_col:
            st.markdown(
                f'<p class="output-meta">'
                f'<strong>{company}</strong> · '
                f'{inputs.get("industry", "")} · {today}'
                f'</p>',
                unsafe_allow_html=True,
            )
        with dl_md_col:
            st.download_button(
                label="Download Markdown",
                data=full_md.encode("utf-8"),
                file_name=f"{slug}_100day_plan_{today}.md",
                mime="text/markdown",
                use_container_width=True,
                key="p_dl_md",
            )
        with dl_docx_col:
            docx_bytes = markdown_to_docx(full_md, company)
            st.download_button(
                label="Download as Word",
                data=docx_bytes,
                file_name=f"{slug}_100day_plan_{today}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="p_dl_docx",
            )
        with dl_pptx_col:
            try:
                from pptx_export import export_pptx  # noqa: E402
                _tpl_bytes = st.session_state.get("pptx_template")
                _tpl_io    = io.BytesIO(_tpl_bytes) if _tpl_bytes else None
                _pptx_buf  = export_pptx(sections, "100day", company, template=_tpl_io)
                st.download_button(
                    label="Export as PowerPoint",
                    data=_pptx_buf.getvalue(),
                    file_name=f"{slug}_100day_plan_{today}.pptx",
                    mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
                    use_container_width=True,
                    key="p_dl_pptx",
                )
            except Exception as _e:
                st.caption(f"PPT unavailable: {_e}")

        st.markdown("<br>", unsafe_allow_html=True)

        # Section expanders — workstreams open by default
        for key, title in _PLAN_SECTION_TITLES.items():
            content = sections.get(key, "")
            if not content or content.startswith("_Error"):
                continue
            with st.expander(title, expanded=(key == "workstreams")):
                st.markdown('<div class="brief-wrapper">', unsafe_allow_html=True)
                st.markdown(content, unsafe_allow_html=False)
                st.markdown('</div>', unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("Generate another 100-Day Plan", use_container_width=False, key="p_reset"):
            for k in ["p_result", "p_error", "p_inputs", "p_generating"]:
                st.session_state.pop(k, None)
            st.rerun()

    elif not st.session_state.get("p_generating") and not st.session_state.get("p_error"):
        st.markdown("""
<div class="empty-state">
  <div class="icon">📅</div>
  <div class="title">Output will appear here</div>
  <div class="sub">Fill in the deal details on the left and click
    <strong>Generate 100-Day Plan</strong>.</div>
</div>
""", unsafe_allow_html=True)


with tab3:
    col_p_form, col_p_out = st.columns([1, 1.6], gap="large")

    with col_p_form:
        # Pre-fill levers from Tab 1 VCO output (once, before widget renders)
        if "vco_for_100day" in st.session_state and "p_levers_prefilled" not in st.session_state:
            st.session_state["p_value_levers"] = st.session_state["vco_for_100day"]
            st.session_state["p_levers_prefilled"] = True

        st.markdown('<div class="section-label">Target Details</div>', unsafe_allow_html=True)

        p_company = st.text_input(
            "Company name *",
            placeholder="e.g. Acme Packaging",
            key="p_company_name",
        )
        p_industry = st.text_input(
            "Industry vertical *",
            placeholder="e.g. Industrial Packaging, Healthcare Services",
            key="p_industry",
        )
        p_deal_type = st.selectbox(
            "Deal type *",
            options=["Platform", "Tuck-in", "Carve-out"],
            index=0,
            key="p_deal_type",
            help="Sets required workstream categories for this deal structure.",
        )

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Financials</div>', unsafe_allow_html=True)

        fin_col1, fin_col2 = st.columns(2)
        with fin_col1:
            p_entry_ebitda = st.number_input(
                "Entry EBITDA ($M) *",
                min_value=0.0,
                value=None,
                placeholder="e.g. 14.0",
                step=0.5,
                format="%.1f",
                key="p_entry_ebitda",
            )
        with fin_col2:
            p_target_ebitda = st.number_input(
                "Target EBITDA ($M) *",
                min_value=0.0,
                value=None,
                placeholder="e.g. 28.0",
                step=0.5,
                format="%.1f",
                key="p_target_ebitda",
            )

        p_hold = st.slider(
            "Hold period (years)",
            min_value=3,
            max_value=7,
            value=5,
            step=1,
            key="p_hold_period",
        )

        st.markdown("<br>", unsafe_allow_html=True)
        st.markdown('<div class="section-label">Thesis & Context</div>', unsafe_allow_html=True)

        p_thesis = st.text_area(
            "PE investment thesis *",
            placeholder=(
                "Core value creation hypothesis — e.g. Fragmented market with pricing power, "
                "undermanaged working capital, greenfield expansion into adjacent geographies."
            ),
            height=110,
            key="p_thesis",
        )
        p_challenges = st.text_area(
            "Key challenges *",
            placeholder=(
                "Issues surfaced in diligence or known pre-close: no CFO in seat, "
                "ERP is 15 years old, top customer is 38% of revenue, margins compressed."
            ),
            height=100,
            key="p_key_challenges",
        )
        p_mgmt = st.text_area(
            "Management assessment (optional)",
            placeholder=(
                "Known strengths or gaps in the leadership team. "
                "e.g. Strong COO, weak CFO, no CRO — founder still owns all customer relationships."
            ),
            height=90,
            key="p_mgmt_assessment",
        )

        # Top Value Levers — pre-filled from Tab 1 if available
        _levers_label = "Top value levers *"
        if "tab1_run_at" in st.session_state and "p_result" not in st.session_state:
            st.success("Pre-filled from your diligence brief")
        p_levers = st.text_area(
            _levers_label,
            placeholder="One lever per line — e.g.\nPricing increase 3–5% on Product Line A\nDSO reduction from ~52 to ~40 days",
            height=130,
            key="p_value_levers",
            help="One lever per line. These drive workstream selection and the EBITDA bridge.",
        )

        # ── Export Options ────────────────────────────────────────────────────
        st.markdown("<br>", unsafe_allow_html=True)
        with st.expander("Export Options", expanded=False):
            _render_pptx_export_options("p")

        st.markdown("<br>", unsafe_allow_html=True)
        p_generate = st.button(
            "Generate 100-Day Plan",
            type="primary",
            use_container_width=True,
            disabled=st.session_state.get("p_generating", False),
            key="p_generate_btn",
        )

    with col_p_out:
        if p_generate:
            missing = []
            if not p_company.strip():
                missing.append("Company name")
            if not p_industry.strip():
                missing.append("Industry vertical")
            if p_entry_ebitda is None or p_entry_ebitda <= 0:
                missing.append("Entry EBITDA")
            if p_target_ebitda is None or p_target_ebitda <= 0:
                missing.append("Target EBITDA")
            if not p_thesis.strip():
                missing.append("Investment thesis")
            if not p_challenges.strip():
                missing.append("Key challenges")
            if not p_levers.strip():
                missing.append("Top value levers")

            if missing:
                st.error(f"Please fill in: {', '.join(missing)}")
            else:
                levers_list = [l.strip() for l in p_levers.strip().splitlines() if l.strip()]
                st.session_state["p_generating"] = True
                st.session_state["p_result"]     = None
                st.session_state["p_error"]      = None
                st.session_state["p_inputs"]     = {
                    "company_name":    p_company.strip(),
                    "industry":        p_industry.strip(),
                    "deal_type":       p_deal_type,
                    "entry_ebitda":    float(p_entry_ebitda) * 1000,   # convert $M → $K
                    "target_ebitda":   float(p_target_ebitda) * 1000,
                    "hold_period_years": int(p_hold),
                    "pe_thesis":       p_thesis.strip(),
                    "key_challenges":  p_challenges.strip(),
                    "mgmt_assessment": p_mgmt.strip() or "Not provided.",
                    "top_value_levers": levers_list,
                    "model_mode":      st.session_state.get("sb_model_mode", "standard"),
                }
                st.rerun()

        if st.session_state.get("p_generating") and not st.session_state.get("p_result"):
            inputs = st.session_state["p_inputs"]
            from generate_100day import generate_100day_plan  # noqa: E402

            steps = [
                "Building workstream plan (Days 1–30 / 31–60 / 61–100)…",
                "Drafting 3rd party resource plan…",
                "Assessing C-suite and leadership team…",
                "Designing 18-month org structure…",
                "Building EBITDA bridge…",
            ]

            with st.status("Generating 100-Day Plan — typically 30–60 seconds…", expanded=True) as status_box:
                for s in steps:
                    status_box.write(f"⏳ {s}")

                result_q: queue.Queue = queue.Queue()
                t = threading.Thread(
                    target=run_async_in_thread,
                    args=(
                        generate_100day_plan(
                            company_name=inputs["company_name"],
                            industry=inputs["industry"],
                            deal_type=inputs["deal_type"],
                            entry_ebitda=inputs["entry_ebitda"],
                            target_ebitda=inputs["target_ebitda"],
                            hold_period_years=inputs["hold_period_years"],
                            pe_thesis=inputs["pe_thesis"],
                            key_challenges=inputs["key_challenges"],
                            mgmt_assessment=inputs["mgmt_assessment"],
                            top_value_levers=inputs["top_value_levers"],
                            model_mode=inputs.get("model_mode", "standard"),
                        ),
                        result_q,
                    ),
                    daemon=True,
                )
                t.start()
                t.join()

                kind, payload = result_q.get_nowait()

                if kind == "ok":
                    st.session_state["p_result"]     = payload
                    st.session_state["p_generating"] = False
                    status_box.update(label="100-Day Plan ready.", state="complete", expanded=False)
                    st.rerun()
                else:
                    st.session_state["p_error"]      = str(payload)
                    st.session_state["p_generating"] = False
                    status_box.update(label="Generation failed.", state="error", expanded=False)
                    st.rerun()

        _render_100day_output()
